"""
Document Processor Module.
Handles extraction of text content from multiple file formats:
PDF, DOCX, TXT, CSV, and images (via OCR + Vision LLM).
"""

import csv
import io
import base64
import logging
from pathlib import Path
from typing import Optional

from PIL import Image
from PyPDF2 import PdfReader
from docx import Document as DocxDocument

from config import SUPPORTED_EXTENSIONS

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """Unified document processor supporting multiple file formats."""

    def __init__(self, ocr_reader=None):
        """
        Initialize the document processor.

        Args:
            ocr_reader: Pre-initialized EasyOCR reader instance (lazy-loaded if None).
        """
        self._ocr_reader = ocr_reader

    @property
    def ocr_reader(self):
        """Lazy-load EasyOCR reader on first use to avoid slow startup."""
        if self._ocr_reader is None:
            try:
                import easyocr
                self._ocr_reader = easyocr.Reader(["en"], gpu=True)
                logger.info("EasyOCR initialized with GPU support.")
            except Exception:
                import easyocr
                self._ocr_reader = easyocr.Reader(["en"], gpu=False)
                logger.info("EasyOCR initialized (CPU mode).")
        return self._ocr_reader

    def process_file(self, file_name: str, file_bytes: bytes) -> dict:
        """
        Process a file and extract its text content.

        Args:
            file_name: Original filename (used to detect format).
            file_bytes: Raw file content as bytes.

        Returns:
            dict with keys: 'text' (extracted content), 'source' (filename),
            'type' (file type), 'is_image' (bool).

        Raises:
            ValueError: If the file extension is not supported.
        """
        ext = Path(file_name).suffix.lower()

        if ext not in SUPPORTED_EXTENSIONS:
            raise ValueError(
                f"Unsupported file type: '{ext}'. "
                f"Supported: {', '.join(SUPPORTED_EXTENSIONS.keys())}"
            )

        processors = {
            ".pdf": self._process_pdf,
            ".docx": self._process_docx,
            ".txt": self._process_txt,
            ".csv": self._process_csv,
            ".jpg": self._process_image,
            ".jpeg": self._process_image,
            ".png": self._process_image,
        }

        processor = processors[ext]
        text = processor(file_bytes)
        is_image = ext in {".jpg", ".jpeg", ".png"}

        if not text or not text.strip():
            logger.warning(f"No text extracted from '{file_name}'.")
            text = "[No readable text content found in this file.]"

        logger.info(
            f"Processed '{file_name}' ({SUPPORTED_EXTENSIONS[ext]}): "
            f"{len(text)} characters extracted."
        )

        return {
            "text": text.strip(),
            "source": file_name,
            "type": SUPPORTED_EXTENSIONS[ext],
            "is_image": is_image,
        }

    def get_image_base64(self, file_bytes: bytes) -> str:
        """
        Convert image bytes to a base64-encoded string for Vision LLM.

        Args:
            file_bytes: Raw image bytes.

        Returns:
            Base64-encoded string of the image.
        """
        return base64.b64encode(file_bytes).decode("utf-8")

    # ─── Private Processors ─────────────────────────────────────────────

    def _process_pdf(self, file_bytes: bytes) -> str:
        """Extract text from all pages of a PDF."""
        try:
            reader = PdfReader(io.BytesIO(file_bytes))
            pages = []
            for i, page in enumerate(reader.pages):
                page_text = page.extract_text()
                if page_text:
                    pages.append(f"[Page {i + 1}]\n{page_text}")
            return "\n\n".join(pages)
        except Exception as e:
            logger.error(f"PDF processing failed: {e}")
            raise RuntimeError(f"Failed to process PDF: {e}") from e

    def _process_docx(self, file_bytes: bytes) -> str:
        """Extract text from all paragraphs and tables in a DOCX file."""
        try:
            doc = DocxDocument(io.BytesIO(file_bytes))
            parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    parts.append(para.text)

            # Extract tables
            for table in doc.tables:
                table_rows = []
                for row in table.rows:
                    cells = [cell.text.strip() for cell in row.cells]
                    table_rows.append(" | ".join(cells))
                if table_rows:
                    parts.append("\n".join(table_rows))

            return "\n\n".join(parts)
        except Exception as e:
            logger.error(f"DOCX processing failed: {e}")
            raise RuntimeError(f"Failed to process DOCX: {e}") from e

    def _process_txt(self, file_bytes: bytes) -> str:
        """Read plain text files with encoding detection."""
        for encoding in ("utf-8", "utf-8-sig", "latin-1", "cp1252"):
            try:
                return file_bytes.decode(encoding)
            except (UnicodeDecodeError, LookupError):
                continue
        return file_bytes.decode("utf-8", errors="replace")

    def _process_csv(self, file_bytes: bytes) -> str:
        """Convert CSV data into readable text format."""
        try:
            text = self._process_txt(file_bytes)
            reader = csv.reader(io.StringIO(text))
            rows = list(reader)

            if not rows:
                return ""

            # Format as readable table
            header = rows[0]
            lines = [" | ".join(header), "-" * (len(" | ".join(header)))]
            for row in rows[1:]:
                lines.append(" | ".join(row))

            return "\n".join(lines)
        except Exception as e:
            logger.error(f"CSV processing failed: {e}")
            raise RuntimeError(f"Failed to process CSV: {e}") from e

    def _process_image(self, file_bytes: bytes) -> str:
        """Extract text from images using EasyOCR."""
        try:
            image = Image.open(io.BytesIO(file_bytes))
            # Convert to RGB if necessary (e.g., RGBA PNGs)
            if image.mode != "RGB":
                image = image.convert("RGB")

            # Save to buffer for EasyOCR
            buf = io.BytesIO()
            image.save(buf, format="PNG")
            buf.seek(0)

            results = self.ocr_reader.readtext(buf.getvalue())
            extracted_texts = [result[1] for result in results if result[2] > 0.3]

            return " ".join(extracted_texts) if extracted_texts else ""
        except Exception as e:
            logger.error(f"Image OCR failed: {e}")
            raise RuntimeError(f"Failed to process image: {e}") from e
